from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import datetime
import os

SELECTED_FILE = 'selected_news.json'
OUTPUT_FILE = '本周新闻Highlight.docx'

def create_document():
    doc = Document()
    
    # Title
    heading = doc.add_heading('本周新闻 Highlight', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date Range
    today_str = datetime.date.today().strftime("%Y.%m.%d")
    date_para = doc.add_paragraph(f'时间范围: 2026.02.27 - {today_str}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('') # Spacer

    # Load Data
    if not os.path.exists(SELECTED_FILE):
        print(f"Error: {SELECTED_FILE} not found.")
        return

    try:
        with open(SELECTED_FILE, 'r', encoding='utf-8') as f:
            news_items = json.load(f)
    except Exception as e:
        print(f"Error loading json: {e}")
        return

    # Filter items
    filtered_items = [
        item for item in news_items 
        if item.get('date', '') >= '2026-02-27'
    ]
    
    # Sort by date desc
    filtered_items.sort(key=lambda x: x['date'], reverse=True)

    # Pre-defined summaries mapping (ID in link -> points)
    summaries = {
        "p4HFkQLKmIEbBWdNLdBcBA": [ # GPT-5.3 & Gemini 3.1
            "OpenAI 发布 GPT-5.3 Instant，幻觉率降低 **26.8%**，主打更自然的对话风格和细节写作能力，告别“AI腔”。",
            "Google 推出 Gemini 3.1 Flash-Lite，首字响应速度提升 **2.5倍**，支持“思考等级”调节，在 GPQA 测试中准确率达 **86.9%**。",
            "Gemini 3.1 Flash-Lite 定价极具攻击性：输入 **$0.25**/百万 tokens，输出 **$1.50**/百万 tokens，大幅降低了实时应用的成本。"
        ],
        "25864": [ # ChatGPT
            "OpenAI 与美国国防部合作引发用户强烈抵制，导致 ChatGPT 移动端单日卸载量环比激增 **295%**，App Store 一星评价飙升 **775%**。",
            "竞争对手 Anthropic 因拒绝类似合作而受益，其 Claude 模型下载量在同期暴增，并成功登顶美国 App Store 免费榜。",
            "数据表明，用户对 AI 技术军事化的担忧已转化为实际的市场行为，直接影响了产品的留存率和品牌声誉。"
        ],
        "25863": [ # GPT-5.4
            "OpenAI 代码仓库意外泄露尚未发布的 'gpt-5.4' 模型信息，显示其可能具备 **200万** Tokens 的超长上下文窗口。",
            "新模型引入了 '状态化 AI' (Stateful AI) 机制，能够跨会话保留工作流和开发环境，有望彻底改变 AI 协作模式。",
            "此次泄露还暗示了针对图像处理的像素级优化，预示着 OpenAI 在多模态和长文本领域的重大技术突破。"
        ],
        "25861": [ # Qwen
            "阿里通义千问发布 Qwen3.5 小型模型系列，涵盖 **0.8B**、**2B**、**4B**、**9B** 四种尺寸，专为边缘设备设计。",
            "该系列模型在保持极致轻量化的同时，具备原生多模态能力，能够以极低的算力消耗运行复杂的 Agent 任务。",
            "此次开源大幅降低了端侧 AI 的部署门槛，为移动端和 IoT 设备上的实时智能交互提供了强有力的基础模型。"
        ],
        "25856": [ # NotebookLM
            "Google NotebookLM 推出信息图自定义样式功能，提供 **10种** 预设风格及完全自定义选项。",
            "用户仅需 **单击一下**，即可将复杂的文档内容瞬间转化为视觉效果出色的信息图，极大提升了知识分享的效率。",
            "这一更新标志着 AI 笔记工具从单纯的文本整理向多模态内容创作进化，满足了用户对个性化和可视化表达的需求。"
        ],
        "25852": [ # MaxClaw
            "MiniMax 移动应用 MaxClaw 正式全球上线，支持 iOS 和 Android 平台，并引入了灵活的 Coding Plan 计费模式。",
            "为应对发布后的高并发流量，MiniMax 进行了 **4次** 紧急服务扩容，并承诺对受影响用户进行积分补偿。",
            "此举不仅拓展了 MiniMax 的 C 端产品线，也通过移动端的便捷入口加速了其 AI Agent 能力在全球市场的渗透。"
        ],
        "25870": [ # Cursor
            "AI 编程助手 Cursor 年化营收 (ARR) 突破 **20亿美元**，在短短 **三个月** 内实现了翻倍增长。",
            "目前约 **60%** 的收入来自大型企业客户，证明其已成功从个人开发者工具转型为企业级生产力平台。",
            "尽管面临 Claude Code 等低价竞品的挑战，Cursor 凭借在高端市场的深耕和 **2.3亿美元** 的融资支持，依然保持着行业领先地位。"
        ],
        "25888": [ # MiniMax Report
            "MiniMax 发布上市后首份年报，全年收入达到 **7904万美元**，同比增长 **158.9%**，其中海外收入占比超过 **70%**。",
            "尽管因金融负债公允价值重估导致账面亏损达 **18.72亿美元**，但销售费用大幅下降，经营效率显著提升。",
            "财报数据显示，MiniMax 正通过高性价比模型和全球化 C 端应用，走出一条不同于传统大模型厂商的商业化路径。"
        ],
        "25877": [ # Claude Voice
            "Anthropic 为 Claude Code 推出官方语音模式，允许开发者通过自然语言指令进行编程，目前已向 **5%** 的用户推送。",
            "用户只需输入 `/voice` 即可开启功能，支持实时语音转录和混合输入，特别适合快速描述复杂逻辑或重构代码。",
            "这一功能的加入进一步降低了编程门槛，将人机协作从键盘输入扩展到语音交互，提升了开发效率和体验。"
        ],
        "25908": [ # MiniMax 2.5
            "MiniMax M2.5 发布仅一周，全球调用量即突破 **3.07万亿** Tokens，月收入 (ARR) 飙升至 **1.5亿美元**。",
            "OpenRouter 数据显示，MiniMax M2.5、Kimi K2.5 和 GLM-5 包揽了全球调用量前三名，展现了国产大模型的统治力。",
            "M2.5 凭借 **10B** 的小参数和极致性价比，精准切中 Agent 市场痛点，证明了效率和成本是当前 AI 应用落地的关键。"
        ]
    }

    print(f"Found {len(filtered_items)} items to process.")

    for idx, item in enumerate(filtered_items, 1):
        # Identify key from link
        key = None
        for k in summaries:
            if k in item['link']:
                key = k
                break
        
        # Heading
        p_title = doc.add_paragraph()
        run_title = p_title.add_run(f"{idx}. {item['title']}")
        run_title.bold = True
        run_title.font.size = Pt(14)
        run_title.font.color.rgb = RGBColor(0, 0, 0)
        
        # Source & Date
        p_meta = doc.add_paragraph()
        run_meta = p_meta.add_run(f"来源: {item['source']} | 日期: {item['date']}")
        run_meta.italic = True
        run_meta.font.size = Pt(10)
        run_meta.font.color.rgb = RGBColor(100, 100, 100)
        
        # Summary Points
        points = []
        if key:
            points = summaries[key]
        else:
            # Fallback to item summary split by sentences or just the full text
            text = item.get('summary', '暂无总结')
            points = [text]

        for point in points:
            p = doc.add_paragraph(style='List Bullet')
            # Handle **bolding**
            parts = point.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1: # Odd parts are between **
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 51, 153) # Dark Blue for data
        
        doc.add_paragraph('') # Spacer

    doc.save(OUTPUT_FILE)
    print(f"Document saved to {OUTPUT_FILE}")

if __name__ == "__main__":
    create_document()
